#!/usr/bin/env python
# -*- coding: utf-8 -*-


from docx import Document
from docx.shared import Inches

from Scrapper import ChemieScrapper


class ChemieDocumentCreator9000:

    def __init__(self, driver=None, delays=None, append:bool = False):

        self.scrapper = None

        if driver is not None:
            self.scrapper = ChemieScrapper(driver=driver, delays=delays)
        else:
            self.scrapper = ChemieScrapper(delays=delays)

        self.delays = None
        if delays is not None:
            self.delays = delays

        self.document_name = 'Sicherheitstabelle.docx'

        if append:
            self.document = Document(self.document_name)
        else:
            self.document = Document()
        # self.table_header = ["Substanz", "Gefahrenpiktrogramm", "H-Sätze", "P-Sätze", "Entsorgung", "CAS-Nr.", "Mol Masse", "Eingesetzte Menge"]
        self.table_header = ["Stoff", "CAS-Nr.", "M in g/mol", "Gefahrensymbol und Signalwort", "H-Sätze", "P-Sätze", "Entsorgung", "Eingesetzte Menge"]

        self.not_found_stoffe = []

        self.content = []

    def extract_table(self, zvg:str=None):

        hs, ps, sw, srcs, name, gefahr, cas_number, mol_masse = self.scrapper.scrape(zvg)

        if len(hs) == 0 and len(ps) == 0 and len(srcs) == 0:
            sw = gefahr

        self.content.append([hs, ps, sw, srcs, name, cas_number, mol_masse])



    def save(self, name=None):

        #######################################
        #         TEXT BEFORE THE TABLE       #
        #######################################
        paragraph_before_table = self.document.add_paragraph()
        paragraph_before_table = paragraph_before_table.add_run()
        paragraph_before_table.add_text("")

        # create table
        table = self.document.add_table(rows=1, cols=8)
        table.style = 'TableGrid'
        for i in range(8):
            table.rows[0].cells[i].text = self.table_header[i]

        for hs, ps, sw, srcs, name, cas_number, mol_masse in self.content:
            data = [
                [name, "", ", ".join([h[0] for h in hs]), ", ".join([h[0] for h in ps]), "(1)", cas_number, mol_masse]
            ]

            for su, ge, hs, ps, en, cas_number, mol_masse in data:
                cells = table.add_row().cells
                cells[0].text = su
                cells[1].text = cas_number
                cells[2].text = mol_masse
                cells[3].text = ge
                p = cells[3].add_paragraph()
                r = p.add_run()
                for alt in srcs:
                    r.add_picture(f'./images/{alt}.gif', width=Inches(0.472441))
                r.add_text("")
                r.add_text(sw)

                cells[4].text = hs
                cells[5].text = ps
                cells[6].text = en

        # add not found elements
        for name in self.not_found_stoffe:
            cells = table.add_row().cells
            cells[0].text = name
            cells[1].text = "Stoff wurde nicht in der GESTIS-Datenbank gefunden."



        #######################################
                # TEXT AFTER THE TABLE
        #######################################

        paragraph_after_table = self.document.add_paragraph()
        paragraph_after_table = paragraph_after_table.add_run()
        paragraph_after_table.add_text("Wortlaut der oben genanten H- und P-Sätze")

        table = self.document.add_table(rows=1, cols=2)
        table.style = 'TableGrid'
        HHH = []
        PPP = []
        for hs, ps, sw, srcs, name, cas_number, mol_masse in self.content:
            # data = [
            #     [name, "", ", ".join(hs[0]), ", ".join(ps[0]), "(1)", cas_number, mol_masse]
            # ]
            # for su, ge, H, P, en, cas_number, mol_masse in data:
            # print(hs)
            for stoffe, text in hs:
                # stoff = stoffe
                HHH.append(f'{stoffe}: {text}')
            # print(ps)
            for stoffe, text in ps:
                # stoff = ", ".join(stoffe)
                # stoff, text = ps
                PPP.append(f'{stoffe}: {text}')

        table.rows[0].cells[0].text = '\n'.join(HHH)
        table.rows[0].cells[1].text = '\n'.join(PPP)

        # mit inhalt
        paragraph_after_table = self.document.add_paragraph()
        paragraph_after_table = paragraph_after_table.add_run()
        paragraph_after_table.add_text("")

        self.document.save(self.document_name)

if __name__ == '__main__':
    extractor = ChemieDocumentCreator9000()
    extractor.extract_table('008230')
    extractor.save()
